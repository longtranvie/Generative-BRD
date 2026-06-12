"""Build a .docx version of the BRD from the structured draft JSON.

Works from draft_N_json rather than the rendered markdown so no markdown
parsing is needed — the draft's sections carry name/content/sources/emphasis
directly (same shape the renderer node consumes).
"""
import io
import re
from datetime import datetime, timezone

from docx import Document
from docx.shared import Pt

_BULLET_RE = re.compile(r"^[-*]\s+")
_NUMBER_RE = re.compile(r"^\d+[.)]\s+")
# **bold** or *italic* tokens; no nesting support — good enough for LLM prose.
_INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*[^*\s][^*]*\*)")


def _add_runs(paragraph, text: str) -> None:
    for token in _INLINE_RE.split(text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**") and len(token) > 4:
            paragraph.add_run(token[2:-2]).bold = True
        elif token.startswith("*") and token.endswith("*") and len(token) > 2:
            paragraph.add_run(token[1:-1]).italic = True
        else:
            paragraph.add_run(token)


def _add_meta_line(doc, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)


def build_brd_docx(values: dict) -> bytes:
    """Render the latest draft (draft_2_json, else draft_1_json) to docx bytes.

    Raises ValueError when the session has no draft yet — callers map this
    to an HTTP 404.
    """
    draft = values.get("draft_2_json") or values.get("draft_1_json")
    if not draft or not draft.get("sections"):
        raise ValueError("No draft to export yet.")

    doc = Document()
    doc.add_heading("Business Requirements Document", level=0)

    meta_bits = []
    if values.get("source_filename"):
        meta_bits.append(f"Source: {values['source_filename']}")
    if values.get("attempt_number"):
        meta_bits.append(f"Attempt {values['attempt_number']}")
    if values.get("status"):
        meta_bits.append(str(values["status"]))
    meta_bits.append(datetime.now(timezone.utc).strftime("Exported %Y-%m-%d %H:%M UTC"))
    _add_meta_line(doc, " · ".join(meta_bits))

    for section in draft.get("sections", []):
        doc.add_heading(section.get("name") or "Untitled section", level=1)
        emphasis = section.get("emphasis")
        if emphasis:
            _add_meta_line(doc, f"Emphasis: {str(emphasis).replace('_', ' ').upper()}")
        for line in (section.get("content") or "").splitlines():
            line = line.strip()
            if not line:
                continue
            if _BULLET_RE.match(line):
                p = doc.add_paragraph(style="List Bullet")
                _add_runs(p, _BULLET_RE.sub("", line))
            elif _NUMBER_RE.match(line):
                p = doc.add_paragraph(style="List Number")
                _add_runs(p, _NUMBER_RE.sub("", line))
            else:
                p = doc.add_paragraph()
                _add_runs(p, line)
        sources = section.get("sources") or []
        if sources:
            _add_meta_line(doc, "Sources: " + ", ".join(sources))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
