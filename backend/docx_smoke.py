"""Smoke for export_docx.build_brd_docx — run directly, no server needed."""
import io

from docx import Document

from export_docx import build_brd_docx

FIXTURE = {
    "source_filename": "discovery.md",
    "attempt_number": 2,
    "status": "FINAL",
    "draft_2_json": {
        "sections": [
            {
                "name": "Executive Summary",
                "emphasis": "must_have",
                "content": (
                    "The portal cuts onboarding from **6 weeks** to *3 days*.\n"
                    "\n"
                    "- SSO self-serve\n"
                    "- First dashboard wizard\n"
                    "1. Numbered point"
                ),
                "sources": ["chunk_1", "chunk_7"],
            },
            {"name": "Empty Section", "content": ""},
        ]
    },
}


def main() -> None:
    data = build_brd_docx(FIXTURE)
    assert data[:2] == b"PK", "docx must be a zip (PK magic)"

    doc = Document(io.BytesIO(data))
    texts = [p.text for p in doc.paragraphs]
    styles = [p.style.name for p in doc.paragraphs]

    assert "Business Requirements Document" in texts
    assert "Executive Summary" in texts
    assert "Empty Section" in texts, "empty sections keep their heading"
    assert any("Sources: chunk_1, chunk_7" in t for t in texts)
    assert "List Bullet" in styles, "bullet lines must use the bullet style"
    assert "List Number" in styles, "numbered lines must use the number style"

    bold_runs = [r.text for p in doc.paragraphs for r in p.runs if r.bold]
    italic_runs = [r.text for p in doc.paragraphs for r in p.runs if r.italic]
    assert "6 weeks" in bold_runs, f"**bold** must become a bold run, got {bold_runs}"
    assert "3 days" in italic_runs, f"*italic* must become an italic run, got {italic_runs}"

    # A draftless session must raise — the endpoint turns this into a 404.
    try:
        build_brd_docx({"draft_1_json": None, "draft_2_json": None})
        raise AssertionError("expected ValueError for missing draft")
    except ValueError:
        pass

    print("DOCX SMOKE PASSED")


if __name__ == "__main__":
    main()
